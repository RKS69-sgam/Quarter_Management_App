import streamlit as st
import pandas as pd
import math
from datetime import datetime, timedelta
import os
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
import io

# --- 1. CONFIG & SECURITY ---
st.set_page_config(page_title="Railway Promotion System", layout="wide")

def check_password():
    if "password_correct" not in st.session_state:
        st.title("🔐 SGAM Office Secure Login")
        user = st.text_input("Username", key="login_user")
        pwd = st.text_input("Password", type="password", key="login_pwd")
        if st.button("Login"):
            if user == "admin" and pwd == "sgam123":
                st.session_state["password_correct"] = True
                st.rerun()
            else:
                st.error("Invalid Credentials")
        return False
    return True

# --- 2. PAY MATRIX DATA ---
PAY_LEVEL_MAP = {
    "1": {"PB": "5200-20200", "GP": "1800"},
    "2": {"PB": "5200-20200", "GP": "1900"},
    "3": {"PB": "5200-20200", "GP": "2000"},
    "4": {"PB": "5200-20200", "GP": "2400"},
    "5": {"PB": "5200-20200", "GP": "2800"},
    "6": {"PB": "9300-34800", "GP": "4200"},
    "7": {"PB": "9300-34800", "GP": "4600"},
    "8": {"PB": "9300-34800", "GP": "4800"},
}

PAY_MATRIX = {
    "1": [18000, 18500, 19100, 19700, 20300, 20900, 21500, 22100, 22800, 23500, 24200, 24900, 25600],
    "2": [19900, 20500, 21100, 21700, 22400, 23100, 23800, 24500, 25200, 26000, 26800, 27600, 28400],
    "3": [21700, 22400, 23100, 23800, 24500, 25200, 26000, 26800, 27600, 28400, 29300, 30200, 31100],
    "4": [25500, 26300, 27100, 27900, 28700, 29600, 30500, 31400, 32300, 33300, 34300, 35300, 36400],
    "5": [29200, 30100, 31000, 31900, 32900, 33900, 34900, 35900, 37000, 38100, 39200, 40400, 41600],
    "6": [35400, 36500, 37600, 38700, 39900, 41100, 42300, 43600, 44900, 46200, 47600, 49000, 50500],
    "7": [44900, 46200, 47600, 49000, 50500, 52000, 53600, 55200, 56900, 58600, 60400, 62200, 64100],
    "8": [47600, 49000, 50500, 52000, 53600, 55200, 56900, 58600, 60400, 62200, 64100, 66000, 68000],
}

# --- 3. HELPERS ---
def find_details(level, pay):
    cells = PAY_MATRIX.get(str(level), [])
    for val in cells:
        if val >= pay: return int(val), cells.index(val) + 1
    return int(pay), len(cells)

def powerful_replace(doc, data):
    for k, v in data.items():
        tag = f"[{k}]"
        for p in doc.paragraphs:
            if tag in p.text: p.text = p.text.replace(tag, str(v))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if tag in p.text: p.text = p.text.replace(tag, str(v))

@st.cache_resource
def init_db():
    if not firebase_admin._apps:
        if "firebase_config" in st.secrets:
            cred_dict = dict(st.secrets["firebase_config"])
            cred_dict["private_key"] = cred_dict["private_key"].replace("\\n", "\n")
            cred = credentials.Certificate(cred_dict)
        else:
            cred = credentials.Certificate('sgamoffice-firebase-adminsdk-fbsvc-253915b05b.json')
        firebase_admin.initialize_app(cred)
    return firestore.client()

# --- 4. MAIN ---
if check_password():
    db = init_db()
    tab1, tab2 = st.tabs(["🚀 Promotion Entry", "📊 Promotion History Report"])

    docs = db.collection("employees").stream()
    df_emp = pd.DataFrame([d.to_dict() | {"id": d.id} for d in docs])
    
    with tab1:
        if not df_emp.empty:
            search_list = df_emp.apply(lambda r: f"{r['Employee Name']} ({str(r.get('PF Number','')).split('.')[0]})", axis=1).tolist()
            sel_emp_str = st.selectbox("Search Employee", search_list)
            emp_data = df_emp[df_emp['Employee Name'] == sel_emp_str.split(' (')[0]].iloc[0]
            
            promo_option = st.radio("Choose Option", ["On Date Promotion", "Promotion On Next Increment"], horizontal=True)
            
            with st.form("final_promo_form"):
                old_pay = int(float(emp_data.get('BASIC PAY', 18000)))
                old_lvl = str(int(float(emp_data.get('PAY LEVEL', 1))))
                
                c1, c2, c3 = st.columns(3)
                curr_pay = c1.number_input("Current Basic Pay", value=old_pay)
                fix_date = c2.date_input("Fixation Date", value=datetime.now())
                order_no = st.text_input("Order Number")

                col1, col2 = st.columns(2)
                with col1:
                    old_desig = emp_data.get('Designation', '')
                    old_gp = PAY_LEVEL_MAP.get(old_lvl, {}).get("GP", "1800")
                    _, old_idx = find_details(old_lvl, curr_pay)
                    st.info(f"Old: {old_desig} | GP: {old_gp}")
                
                with col2:
                    all_eng = sorted(list(set(df_emp['Designation'].dropna())))
                    all_hindi = sorted(list(set(df_emp['Designation in Hindi'].dropna())))
                    new_lvl = st.selectbox("New Level", list(PAY_LEVEL_MAP.keys()), index=int(old_lvl)-1 if int(old_lvl)<8 else 7)
                    new_gp = PAY_LEVEL_MAP[new_lvl]["GP"]
                    new_desig_eng = st.selectbox("New Designation (English)", all_eng)
                    new_desig_hindi = st.selectbox("New Designation (Hindi)", all_hindi)

                # Fixation Calculations
                notional_general = math.ceil((curr_pay * 1.03) / 100) * 100
                general_final_pay, general_new_idx = find_details(new_lvl, notional_general)
                
                val_in_new_gp, initial_new_idx = find_details(new_lvl, curr_pay)
                double_inc_val = math.ceil((curr_pay * 1.03 * 1.03) / 100) * 100
                final_fixation_val, revised_new_idx = find_details(new_lvl, double_inc_val)
                
                inc_year = fix_date.year + (1 if fix_date.month > 6 else 0)
                inc_date_obj = datetime.strptime(f"01.07.{inc_year}" if fix_date.month <= 6 else f"01.01.{inc_year+1}", "%d.%m.%Y")
                prev_day_inc = (inc_date_obj - timedelta(days=1)).strftime("%d.%m.%Y")

                if st.form_submit_button("Process Promotion"):
                    # Update Logic
                    final_db_pay = final_fixation_val if promo_option == "Promotion On Next Increment" else general_final_pay
                    
                    db.collection("employees").document(emp_data['id']).update({
                        "BASIC PAY": int(final_db_pay),
                        "PAY LEVEL": new_lvl,
                        "Designation": new_desig_eng,
                        "Designation in Hindi": new_desig_hindi,
                        "Posting Status": new_desig_eng
                    })
                    
                    # History
                    db.collection("promotion_history").add({
                        "Name": emp_data['Employee Name'],
                        "PF": str(emp_data.get('PF Number','')).split('.')[0],
                        "Type": promo_option,
                        "GP": new_gp,
                        "timestamp": datetime.now()
                    })

                    mapping = {
                        "PFNUMBER": str(emp_data.get('PF Number','')).split('.')[0],
                        "EMPLOYEENAME": emp_data.get('Employee Name in Hindi', emp_data['Employee Name']),
                        "STATION": emp_data.get('STATION', 'N/A'),
                        "OLDDESIGNATION": old_desig, "NEWDESIGNATION": new_desig_hindi,
                        "OLDGP": old_gp, "NEWGP": new_gp,
                        "OLDBASICPAY": int(curr_pay), 
                        "NEWBASICPAY": general_final_pay if promo_option == "On Date Promotion" else val_in_new_gp,
                        "OLDLEVEL": old_lvl, "NEWLEVEL": new_lvl,
                        "OLDPAYBAND": PAY_LEVEL_MAP[old_lvl]["PB"], "NEWPAYBAND": PAY_LEVEL_MAP[new_lvl]["PB"],
                        "OLDINDEX": old_idx, "NEWINDEX": initial_new_idx if promo_option == "Promotion On Next Increment" else general_new_idx,
                        "REVISEDNEWINDEX": revised_new_idx,
                        "PROMOTIONORDERNUMBER": order_no,
                        "PROMOTIONDATE": fix_date.strftime("%d.%m.%Y"),
                        "NEXTINCRDATE": inc_date_obj.strftime("%d.%m.%Y"),
                        "NEXTINCRDATE-1": prev_day_inc,
                        "OLDBASICKPAY=<INNEWGP": val_in_new_gp,
                        "MROUND100OLDBASICPAY*103%": int(notional_general),
                        "MROUND100OLDBASICPAY*103%*103%": double_inc_val,
                        "MROUND100OLDBASICPAY*103%*103%=<INNEWGP": final_fixation_val
                    }

                    t_name = "General Promotion MACP temp.docx" if promo_option == "On Date Promotion" else "On Increment Promotion temp.docx"
                    t_path = os.path.join("assets", t_name)
                    
                    if os.path.exists(t_path):
                        doc = Document(t_path)
                        powerful_replace(doc, mapping)
                        bio = io.BytesIO()
                        doc.save(bio)
                        safe_name = "".join([c if c.isalnum() else "_" for c in emp_data['Employee Name']])
                        st.session_state.file_output = bio.getvalue()
                        st.session_state.file_name = f"{safe_name}_{new_gp}.docx"
                        st.success("Database Updated & Document Ready!")
                        st.rerun()

        if "file_output" in st.session_state:
            st.download_button(f"📥 Download {st.session_state.file_name}", st.session_state.file_output, st.session_state.file_name)

    with tab2:
        st.header("📋 Promotion History Report")
        h_docs = db.collection("promotion_history").order_by("timestamp", direction=firestore.Query.DESCENDING).stream()
        h_list = [h.to_dict() for h in h_docs]
        if h_list:
            df_h = pd.DataFrame(h_list)
            if 'timestamp' in df_h.columns:
                df_h['timestamp'] = pd.to_datetime(df_h['timestamp']).dt.tz_localize(None)
            st.dataframe(df_h, use_container_width=True)
            towrite = io.BytesIO()
            df_h.to_excel(towrite, index=False, engine='openpyxl')
            st.download_button("📂 Export to Excel", towrite.getvalue(), "Promotion_History.xlsx")
