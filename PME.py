import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
import io
import os
import firebase_admin
from firebase_admin import credentials, firestore
from dateutil.relativedelta import relativedelta

# --- 0. Path & Firebase Setup ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "assets", "pme memo temp.docx")

@st.cache_resource
def init_db():
    if not firebase_admin._apps:
        try:
            if "firebase_config" in st.secrets:
                cred_dict = dict(st.secrets["firebase_config"])
                if isinstance(cred_dict.get('private_key'), str):
                    cred_dict['private_key'] = cred_dict['private_key'].replace('\\n', '\n')
                cred = credentials.Certificate(cred_dict)
            else:
                cred = credentials.Certificate('sgamoffice-firebase-adminsdk-fbsvc-253915b05b.json')
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Firebase Error: {e}"); return None
    return firestore.client()

db = init_db()

# --- 1. Robust Utilities ---
def get_safe_date(date_val):
    if not date_val or str(date_val).lower() == 'nan' or date_val == "":
        return None
    try:
        # Check if it's already a datetime object (from Firestore)
        if hasattr(date_val, 'to_datetime'):
            return date_val.to_datetime().replace(tzinfo=None)
        # Parse string formats like DD/MM/YYYY
        return pd.to_datetime(date_val, dayfirst=True).to_pydatetime().replace(tzinfo=None)
    except: return None

def calculate_next_pme(last_pme_raw, dob_raw):
    """
    Railway Medical Manual (IRMM) logic:
    Age < 45: Every 4 years
    Age 45-55: Every 2 years
    Age 55+: Every 1 year
    """
    last_pme = get_safe_date(last_pme_raw)
    dob = get_safe_date(dob_raw)
    if not last_pme or not dob:
        return None
    
    age_at_pme = relativedelta(last_pme, dob).years
    if age_at_pme < 45: interval = 4
    elif 45 <= age_at_pme < 55: interval = 2
    else: interval = 1
    
    return last_pme + relativedelta(years=interval)

def calculate_pme_metrics(dob_raw, doa_raw):
    now = datetime.now()
    res = {"age": "N/A", "s_yr": "0", "s_mn": "0", "dob_f": "", "doa_f": ""}
    dt_dob = get_safe_date(dob_raw)
    dt_doa = get_safe_date(doa_raw)
    if dt_dob:
        res["age"] = str(relativedelta(now, dt_dob).years)
        res["dob_f"] = dt_dob.strftime("%d/%m/%Y")
    if dt_doa:
        diff = relativedelta(now, dt_doa)
        res["s_yr"] = str(diff.years)
        res["s_mn"] = str(diff.months)
        res["doa_f"] = dt_doa.strftime("%d/%m/%Y")
    return res

def replace_placeholders(doc, mapping):
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for p in all_paras:
        for key, val in mapping.items():
            placeholder = "{{ " + str(key) + " }}"
            if placeholder in p.text:
                full_text = "".join(run.text for run in p.runs)
                if placeholder in full_text:
                    new_text = full_text.replace(placeholder, str(val))
                    for i, run in enumerate(p.runs):
                        run.text = new_text if i == 0 else ""

# --- 2. Main Interface ---
st.set_page_config(layout="wide", page_title="Railway PME System")

if db is not None:
    docs = db.collection("employees").stream()
    df_emp = pd.DataFrame([{**d.to_dict(), 'id': d.id} for d in docs])

    if 'memo_bytes' not in st.session_state: st.session_state.memo_bytes = None

    # --- 📢 PME DUE ALERT SECTION ---
    if not df_emp.empty:
        st.subheader("⚠️ PME Due Alerts (Next 15 Days)")
        due_list = []
        today = datetime.now()
        alert_window = today + timedelta(days=15)

        for _, row in df_emp.iterrows():
            next_pme = calculate_next_pme(row.get('Last PME'), row.get('DOB'))
            if next_pme:
                # Alert if PME is already overdue or due within 15 days
                if next_pme <= alert_window:
                    status = "OVERDUE" if next_pme < today else "DUE SOON"
                    due_list.append({
                        "Employee": row.get('Employee Name'),
                        "HRMS ID": row.get('HRMS ID'),
                        "Last PME": row.get('Last PME'),
                        "Next PME Due": next_pme.strftime("%d/%m/%Y"),
                        "Status": status
                    })
        
        if due_list:
            due_df = pd.DataFrame(due_list)
            # Styling for alerts
            def color_status(val):
                color = 'red' if val == "OVERDUE" else 'orange'
                return f'background-color: {color}; color: white; font-weight: bold'
            
            st.table(due_df.style.applymap(color_status, subset=['Status']))
        else:
            st.success("✅ Sabhi karmchariyon ki PME up-to-date hai.")
    
    st.divider()

    tab1, tab2, tab3 = st.tabs(["📝 Generate PME Memo", "📊 History", "🛠 Update Database"])

    with tab1:
        if not df_emp.empty:
            emp_names = df_emp.apply(lambda r: f"{r.get('Employee Name')} ({r.get('HRMS ID')})", axis=1).tolist()
            selected = st.selectbox("Select Employee", emp_names)
            h_id = selected.split('(')[-1].strip(')')
            emp_data = df_emp[df_emp['HRMS ID'] == h_id].iloc[0]

            with st.form("pme_gen_form"):
                memo_date = st.date_input("Memo Date", value=datetime.now())
                m = calculate_pme_metrics(emp_data.get('DOB'), emp_data.get('DOA'))
                
                final_mapping = {
                    "dob": m["dob_f"], "doa": m["doa_f"],
                    "name": emp_data.get('Employee Name', ''),
                    "age": m["age"],
                    "father_name": emp_data.get("FATHER'S NAME", ''),
                    "designation": emp_data.get('Designation', ''),
                    "medical_category": emp_data.get('Medical category', ''),
                    "current_date": memo_date.strftime("%d/%m/%Y"),
                    "first_physical_mark": emp_data.get('Physical Mark 1', 'N/A'),
                    "second_physical_mark": emp_data.get('Physical Mark 2', 'N/A'),
                    "last_examined_date": emp_data.get('Last PME', 'N/A'),
                    "last_place": emp_data.get('Last PME Place', 'N/A'),
                    "examiner": emp_data.get('Last Examiner', 'ACMS/NKJ'),
                    "service_year": m["s_yr"], "service_month": m["s_mn"]
                }

                if st.form_submit_button("Generate Memo"):
                    if os.path.exists(TEMPLATE_PATH):
                        doc = Document(TEMPLATE_PATH)
                        replace_placeholders(doc, final_mapping)
                        buf = io.BytesIO()
                        doc.save(buf)
                        st.session_state.memo_bytes = buf.getvalue()
                        db.collection("pme_history").add({**final_mapping, "Timestamp": datetime.now(), "HRMS_ID": h_id})
                        st.success("✅ Document Generated!")
                    else: st.error("Template Not Found!")

            if st.session_state.memo_bytes:
                st.download_button("📥 Download PME Memo", st.session_state.memo_bytes, f"PME_{emp_data.get('Employee Name', '')}.docx")

    with tab2:
        st.header("Recent Generations")
        h_docs = db.collection("pme_history").order_by("Timestamp", direction=firestore.Query.DESCENDING).limit(10).stream()
        h_list = [{**d.to_dict()} for d in h_docs]
        if h_list: st.dataframe(pd.DataFrame(h_list)[['Timestamp', 'name', 'age', 'service_year']], use_container_width=True)

    with tab3:
        st.header("🛠 Update Employee Medical Data")
        if not df_emp.empty:
            t_sel = st.selectbox("Select Employee to Update", emp_names, key="upd_pme_final")
            t_id = t_sel.split('(')[-1].strip(')')
            t_row = df_emp[df_emp['HRMS ID'] == t_id].iloc[0]
            
            with st.form("medical_update_form"):
                col1, col2 = st.columns(2)
                m1 = col1.text_input("Physical Mark 1", t_row.get('Physical Mark 1', ''))
                m2 = col2.text_input("Physical Mark 2", t_row.get('Physical Mark 2', ''))
                lp_date = col1.text_input("Last PME Date (DD/MM/YYYY)", t_row.get('Last PME', ''))
                lp_cat = col2.text_input("Last Medical Category", t_row.get('Medical category', ''))
                lp_place = col1.text_input("Last PME Place", t_row.get('Last PME Place', ''))
                lp_exam = col2.text_input("Last Examiner", t_row.get('Last Examiner', 'ACMS/NKJ'))
                
                if st.form_submit_button("Save & Update Records"):
                    # Validate date before saving
                    try:
                        valid_date = pd.to_datetime(lp_date, dayfirst=True)
                        db.collection("employees").document(t_row['id']).update({
                            "Physical Mark 1": m1, "Physical Mark 2": m2,
                            "Last PME": lp_date, "Medical category": lp_cat,
                            "Last PME Place": lp_place, "Last Examiner": lp_exam
                        })
                        st.success("✅ Employee Data Updated! Alert will disappear if next PME is > 15 days from today.")
                        st.rerun()
                    except:
                        st.error("Invalid Date Format! Use DD/MM/YYYY")
else: st.error("Database Connection Failed.")
