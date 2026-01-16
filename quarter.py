import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import io
import os
import firebase_admin
from firebase_admin import credentials, firestore

# --- 1. Security Login ---
ADMIN_USER = "admin"
ADMIN_PASS = "Sgam@4321"

def check_login():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if not st.session_state.logged_in:
        st.title("🔐 Railway Quarter Management Login")
        with st.form("login_form"):
            user = st.text_input("Username")
            pas = st.text_input("Password", type="password")
            if st.form_submit_button("Login"):
                if user == ADMIN_USER and pas == ADMIN_PASS:
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("❌ Galat Username ya Password")
        return False
    return True

# --- 2. Firebase Setup ---
@st.cache_resource
def init_db():
    if not firebase_admin._apps:
        try:
            if "firebase_config" in st.secrets:
                cred = credentials.Certificate(dict(st.secrets["firebase_config"]))
            else:
                cred = credentials.Certificate('sgamoffice-firebase-adminsdk-fbsvc-253915b05b.json')
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Firebase Error: {e}"); st.stop()
    return firestore.client()

# --- 3. Data Sync & PF Text Formatting ---
def get_master_data():
    db = init_db()
    
    # Employees Fetching (PF Number as Text)
    emp_docs = db.collection("employees").stream()
    emp_map = {}
    for d in emp_docs:
        val = d.to_dict()
        raw_pf = val.get('PF Number', '')
        if raw_pf:
            # .0 hatane aur text format maintain karne ke liye logic
            pf_str = str(raw_pf).split('.')[0].strip()
            emp_map[pf_str] = val
    
    # History Fetching
    hist_docs = db.collection("quarter_history").stream()
    data = []
    for d in hist_docs:
        item = d.to_dict()
        item['id'] = d.id
        
        # History PF Number cleaning
        raw_h_pf = item.get('pf_number', '')
        item['pf_number'] = str(raw_h_pf).split('.')[0].strip() if raw_h_pf else ""

        # Live Sync from Employees
        curr_pf = item['pf_number']
        if curr_pf in emp_map:
            item['designation'] = emp_map[curr_pf].get('Designation', item.get('designation', ''))
            item['unit'] = emp_map[curr_pf].get('Unit', item.get('unit', ''))
            if not item.get('station'):
                item['station'] = emp_map[curr_pf].get('STATION', 'N/A')
        
        # Display Format logic
        item['allot_disp'] = item['allotment_date'].strftime('%d-%m-%Y') if item.get('allotment_date') and hasattr(item['allotment_date'], 'strftime') else "N/A"
        item['vacat_disp'] = item['vacation_date'].strftime('%d-%m-%Y') if item.get('vacation_date') and hasattr(item['vacation_date'], 'strftime') else ("Occupied" if item.get('is_current') else "Vacant")
        item['status_disp'] = "🔴 Occupied" if item.get('is_current') else "🟢 Vacant"
        
        data.append(item)
        
    return pd.DataFrame(data), emp_map

# --- 4. Template Generator ---
def fill_template(temp_path, data, date_str):
    if not os.path.exists(temp_path):
        st.error(f"Template File Missing: {temp_path}")
        return None
    doc = Document(temp_path)
    mapping = {
        "EMPLOYEE_NAME": str(data.get('employee_name', '')),
        "DESIGNATION": str(data.get('designation', '')),
        "PF_Number": str(data.get('pf_number', '')),
        "HRMS_ID": str(data.get('hrms_id', '')),
        "QUARTER_NUMBER": str(data.get('quarter_number', '')),
        "STATION": str(data.get('station', '')),
        "UNIT": str(data.get('unit', '')),
        "DATE": date_str
    }
    for p in list(doc.paragraphs) + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]:
        for k, v in mapping.items():
            if f"{{{{{k}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{k}}}}}", v)
    return doc

# --- 5. Main Application ---
def main():
    if not check_login(): return

    st.set_page_config(layout="wide", page_title="Railway Quarter MS")
    db = init_db()
    
    if st.sidebar.button("🚪 Logout"):
        st.session_state.logged_in = False
        st.rerun()

    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    ALLOT_TEMP = os.path.join(BASE_DIR, "assets", "Quarter_Allotment_Template.docx")
    VACATE_TEMP = os.path.join(BASE_DIR, "assets", "Quarter_Vacation_Template.docx")

    df_hist, emp_map = get_master_data()
    
    tab1, tab2, tab3 = st.tabs(["🏠 Allotment", "🗝️ Vacation", "📊 Dashboard & History"])

    # --- TAB 1: ALLOTMENT ---
    with tab1:
        st.header("New Allotment Form")
        if not df_hist.empty and emp_map:
            # Alphanumeric PF numbers display logic
            emp_list = sorted([f"{v.get('Employee Name')} ({k})" for k, v in emp_map.items()])
            selected_emp = st.selectbox("Staff Chunein", emp_list)
            
            pf_key = selected_emp.split('(')[-1].strip(')')
            
            # Double Allotment Check
            already_occ = df_hist[(df_hist['pf_number'] == pf_key) & (df_hist['is_current'] == True)]
            
            if not already_occ.empty:
                st.error(f"❌ Is karmchari ko pehle se Quarter No. {already_occ.iloc[0]['quarter_number']} allot hai.")
            else:
                staff_station = str(emp_map[pf_key].get('STATION', '')).strip()
                st.info(f"📍 Station: **{staff_station}**")

                # Filter Vacant Quarters ONLY at this Station
                stn_data = df_hist[df_hist['station'].str.strip() == staff_station]
                occupied_qs = stn_data[stn_data['is_current'] == True]['quarter_number'].unique().tolist()
                vacant_qs = [q for q in sorted(stn_data['quarter_number'].unique()) if q not in occupied_qs]

                if vacant_qs:
                    sel_q = st.selectbox(f"Available Quarters in {staff_station}", vacant_qs)
                    a_date = st.date_input("Allotment Effective Date", value=datetime.now())

                    if st.button("Process Allotment & Generate Letter"):
                        emp_info = emp_map[pf_key]
                        entry = {
                            "employee_name": emp_info.get('Employee Name'),
                            "designation": emp_info.get('Designation'),
                            "unit": emp_info.get('Unit', ''),
                            "hrms_id": emp_info.get('HRMS ID', ''),
                            "pf_number": pf_key,
                            "quarter_number": sel_q,
                            "station": staff_station,
                            "allotment_date": datetime.combine(a_date, datetime.min.time()),
                            "is_current": True
                        }
                        db.collection("quarter_history").add(entry)
                        st.success(f"✅ Quarter {sel_q} allot kar diya gaya.")
                        
                        doc = fill_template(ALLOT_TEMP, entry, a_date.strftime("%d/%m/%Y"))
                        if doc:
                            buf = io.BytesIO(); doc.save(buf)
                            st.download_button("📥 Download Allotment Letter", buf.getvalue(), f"Allotment_{emp_info.get('Employee Name')}.docx")
                else:
                    st.warning(f"⚠️ {staff_station} station par koi vacant quarter nahi mila.")

    # --- TAB 2: VACATION ---
    with tab2:
        st.header("Process Quarter Vacation")
        occ_df = df_hist[df_hist['is_current'] == True]
        if not occ_df.empty:
            v_list = occ_df.apply(lambda r: f"{r['station']} | {r['quarter_number']} - {r['employee_name']}", axis=1).tolist()
            sel_v = st.selectbox("Select Allotted Quarter", v_list)
            v_date = st.date_input("Vacation Effective Date")
            
            if st.button("Vacate & Generate Letter"):
                idx = v_list.index(sel_v)
                q_row = occ_df.iloc[idx]
                
                db.collection("quarter_history").document(q_row['id']).update({
                    "is_current": False,
                    "vacation_date": datetime.combine(v_date, datetime.min.time())
                })
                
                doc = fill_template(VACATE_TEMP, q_row.to_dict(), v_date.strftime("%d/%m/%Y"))
                if doc:
                    buf = io.BytesIO(); doc.save(buf)
                    st.success(f"✅ Quarter {q_row['quarter_number']} Vacated.")
                    st.download_button("📥 Download Vacation Letter", buf.getvalue(), f"Vacation_{sel_v}.docx")
        else:
            st.info("Abhi koi quarter occupied nahi hai.")

    # --- TAB 3: DASHBOARD & HISTORY ---
    with tab3:
        st.header("📊 Station-Wise Summary & History")
        if not df_hist.empty:
            # Correct Counting: Combined Station + Quarter (Alphanumeric safe)
            df_hist['unique_key'] = df_hist['station'].astype(str).str.strip() + "_" + df_hist['quarter_number'].astype(str).str.strip()
            
            total_unique = df_hist['unique_key'].nunique()
            total_occ = df_hist[df_hist['is_current'] == True]['unique_key'].nunique()
            total_vac = total_unique - total_occ
            
            c1, c2, c3 = st.columns(3)
            c1.metric("🏠 Total Quarters (Unique)", total_unique)
            c2.metric("🔴 Currently Occupied", total_occ)
            c3.metric("🟢 Currently Vacant", total_vac)

            

            st.markdown("---")
            st.subheader("🔍 Quarter History Search")
            h1, h2 = st.columns(2)
            with h1:
                s_stn = st.selectbox("Station Chunein", sorted(df_hist['station'].unique()))
            with h2:
                s_q = st.selectbox("Quarter No. Chunein", sorted(df_hist[df_hist['station'] == s_stn]['quarter_number'].unique()))
            
            q_history_data = df_hist[(df_hist['station'] == s_stn) & (df_hist['quarter_number'] == s_q)].sort_values(by='allotment_date', ascending=False)
            
            if not q_history_data.empty:
                hist_disp = q_history_data[['employee_name', 'pf_number', 'allot_disp', 'vacat_disp', 'status_disp']]
                hist_disp.columns = ['Occupied By', 'PF Number', 'Allotment Date', 'Vacation Date', 'Current Status']
                st.table(hist_disp)

            st.markdown("---")
            st.subheader("📋 Master Quarter Report")
            report_cols = ['station', 'quarter_number', 'employee_name', 'pf_number', 'allot_disp', 'status_disp']
            st.dataframe(df_hist[report_cols], use_container_width=True)
            
            csv = df_hist[report_cols].to_csv(index=False).encode('utf-8-sig')
            st.download_button("📥 Download Report CSV", csv, "Quarter_Master_Report.csv", "text/csv")

if __name__ == "__main__":
    main()
