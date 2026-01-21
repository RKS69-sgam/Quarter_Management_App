import streamlit as st
import pandas as pd
import math
from datetime import datetime
import os
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
import io

# =================================================================
# --- 0. CONFIG & AUTHENTICATION ---
# =================================================================
st.set_page_config(page_title="Railway Promotion System", layout="wide")

if 'auth' not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    st.title("🔒 Admin Login")
    with st.form("login"):
        u = st.text_input("User")
        p = st.text_input("Password", type="password")
        if st.form_submit_button("Login"):
            if u == "admin" and p == "Sgam@4321":
                st.session_state.auth = True
                st.rerun()
            else:
                st.error("Invalid credentials")
    st.stop()

# =================================================================
# --- 1. FIREBASE CONNECTION (Secrets & Local Support) ---
# =================================================================
@st.cache_resource
def init_db():
    if not firebase_admin._apps:
        try:
            # Check Streamlit Secrets first
            if "firebase_config" in st.secrets:
                cred_dict = dict(st.secrets["firebase_config"])
                cred_dict["private_key"] = cred_dict["private_key"].replace("\\n", "\n")
                cred = credentials.Certificate(cred_dict)
            else:
                # Fallback to local file
                cred = credentials.Certificate('sgamoffice-firebase-adminsdk-fbsvc-253915b05b.json')
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Database Connection Error: {e}")
            st.stop()
    return firestore.client()

db = init_db()

# =================================================================
# --- 2. PAY MATRIX & LOGIC ---
# =================================================================
PAY_MATRIX = {
    "1": [18000, 18500, 19100, 19700, 20300, 20900, 21500, 22100, 22800],
    "2": [19900, 20500, 21100, 21700, 22400, 23100, 23800, 24500, 25200],
    "3": [21700, 22400, 23100, 23800, 24500, 25200, 26000, 26800, 27600],
    "4": [25500, 26300, 27100, 27900, 28700, 29600, 30500, 31400, 32300],
    "5": [29200, 30100, 31000, 31900, 32900, 33900, 34900, 35900, 37000],
    "6": [35400, 36500, 37600, 38700, 39900, 41100, 42300, 43600, 44900],
    "7": [44900, 46200, 47600, 49000, 50500, 52000, 53600, 55200, 56900]
}

def get_next_increment_date(promo_date):
    return f"01/01/{promo_date.year + 1}" if promo_date.month <= 6 else f"01/07/{promo_date.year + 1}"

def find_cell_in_level(level, target_val):
    cells = PAY_MATRIX.get(str(level), [])
    for val in cells:
        if val >= target_val:
            idx = cells.index(val) + 1
            return val, idx
    return target_val, 1

def generate_docx(template_path, data):
    if not os.path.exists(template_path): return None
    doc = Document(template_path)
    # [span_2](start_span)[span_3](start_span)Search in paragraphs and tables[span_2](end_span)[span_3](end_span)
    for p in list(doc.paragraphs) + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]:
        for k, v in data.items():
            if f"[{k}]" in p.text:
                for run in p.runs:
                    if f"[{k}]" in run.text:
                        run.text = run.text.replace(f"[{k}]", str(v))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =================================================================
# --- 3. UI ---
# =================================================================
tab1, tab2 = st.tabs(["🚀 Promotion Entry", "📜 History Report"])

# Data Loading
emp_docs = db.collection("employees").stream()
df_emp = pd.DataFrame([{**d.to_dict(), 'id': d.id} for d in emp_docs])

with tab1:
    if not df_emp.empty:
        pf_list = df_emp['PF Number'].unique().tolist()
        selected_pf = st.selectbox("Select Employee PF Number", pf_list)
        emp_data = df_emp[df_emp['PF Number'] == selected_pf].iloc[0]
        
        with st.form("promotion_main_form"):
            st.subheader(f"Employee: {emp_data['Employee Name']}")
            col1, col2, col3 = st.columns(3)
            
            # Form Inputs
            old_basic = col1.number_input("Old Basic Pay", value=float(emp_data.get('Basic Pay', 0)))
            promo_date = col2.date_input("Promotion Date", value=datetime.now())
            order_no = col3.text_input("Promotion Order Number")
            
            col4, col5 = st.columns(2)
            new_desig = col4.text_input("New Designation", value=emp_data.get('Designation', ''))
            target_level = col5.selectbox("Target Level", list(PAY_MATRIX.keys()))
            
            # Math Logic
            notional_pay = math.ceil((old_basic * 1.03) / 100) * 100
            final_basic, new_idx = find_cell_in_level(target_level, notional_pay)
            next_date = get_next_increment_date(promo_date)
            
            st.write(f"**New Basic Pay:** ₹{final_basic} (Level {target_level})")
            
            if st.form_submit_button("Update & Generate"):
                # [span_4](start_span)[span_5](start_span)Database Update[span_4](end_span)[span_5](end_span)
                db.collection("employees").document(emp_data['id']).update({
                    "Basic Pay": final_basic,
                    "Level": target_level,
                    "Designation": new_desig
                })
                
                # [span_6](start_span)[span_7](start_span)History Save[span_6](end_span)[span_7](end_span)
                db.collection("promotion_history").add({
                    "PF": selected_pf, "Name": emp_data['Employee Name'], 
                    "NewBasic": final_basic, "Date": str(promo_date), "Timestamp": datetime.now()
                })
                
                # Word Mapping
                mapping = {
                    "PFNUMBER": selected_pf, "EMPLOYEENAME": emp_data.get('Employee Name in Hindi', emp_data['Employee Name']),
                    "OLDBASICPAY": old_basic, "NEWBASICPAY": final_basic, "STATION": emp_data.get('STATION', 'SGAM'),
                    "PROMOTIONDATE": promo_date.strftime("%d.%m.%Y"), "PROMOTIONORDERNUMBER": order_no,
                    "OLDLEVEL": emp_data.get('Level', '1'), "NEWLEVEL": target_level, "NEXTINCRDATE": next_date,
                    "MROUND100OLDBASICPAY*103%": notional_pay
                }
                
                path = os.path.join("assets", "General Promotion MACP temp.docx")
                st.session_state.promo_file = generate_docx(path, mapping)
                st.success("Record updated successfully!")
                st.rerun()

    if 'promo_file' in st.session_state:
        st.download_button("📥 Download Memo", st.session_state.promo_file, f"Promotion_{selected_pf}.docx")

with tab2:
    hist = db.collection("promotion_history").order_by("Timestamp", direction=firestore.Query.DESCENDING).stream()
    st.dataframe(pd.DataFrame([d.to_dict() for d in hist]), use_container_width=True)
