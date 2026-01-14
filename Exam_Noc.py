import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import io
import os
import firebase_admin
from firebase_admin import credentials, firestore

# =================================================================
# --- 0. PATH & FIREBASE SETUP ---
# =================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "assets", "Exam NOC Letter temp.docx")

EMP_COLLECTION = "employees"
NOC_HISTORY_COLLECTION = "noc_history"

@st.cache_resource
def init_db():
    if not firebase_admin._apps:
        try:
            if "firebase_config" in st.secrets:
                cred_dict = dict(st.secrets["firebase_config"])
                if isinstance(cred_dict.get('private_key'), str):
                    cred_dict['private_key'] = cred_dict['private_key'].replace('\\n', '\n')
                cred = credentials.Certificate(cred_dict)
            else:
                cred = credentials.Certificate('sgamoffice-firebase-adminsdk-fbsvc-253915b05b.json')
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Firebase Error: {e}")
            st.stop()
    return firestore.client()

db = init_db()

# =================================================================
# --- 1. UTILITIES ---
# =================================================================
def get_employees():
    docs = db.collection(EMP_COLLECTION).stream()
    data = [{**d.to_dict(), 'id': d.id} for d in docs]
    return pd.DataFrame(data)

def get_noc_history():
    docs = db.collection(NOC_HISTORY_COLLECTION).order_by("Timestamp", direction=firestore.Query.DESCENDING).stream()
    data = [{**d.to_dict(), 'id': d.id} for d in docs]
    return pd.DataFrame(data) if data else pd.DataFrame()

def get_noc_count_this_year(pf_number):
    current_year = datetime.now().year
    docs = db.collection(NOC_HISTORY_COLLECTION)\
             .where("PFNumber", "==", str(pf_number))\
             .where("Year", "==", current_year).stream()
    return len(list(docs))

def create_noc_table(doc, emp_data_list):
    for p in doc.paragraphs:
        if "[PFNumber]" in p.text:
            p.text = p.text.replace("[PFNumber]", "")
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            headers = ['Sr. No.', 'PF Number', 'Employee Name', 'Designation', "Exam's Name", 'Term of NOC']
            for i, h_text in enumerate(headers):
                hdr[i].text = h_text
            
            for idx, emp in enumerate(emp_data_list):
                row = table.add_row().cells
                row[0].text = str(idx + 1)
                row[1].text = str(emp.get('PFNumber', ''))
                row[2].text = str(emp.get('Name', ''))
                row[3].text = str(emp.get('Desig', ''))
                row[4].text = str(emp.get('ExamName', ''))
                row[5].text = str(emp.get('Term', ''))
            break

def generate_multi_noc(template_path, l_date, emp_data_list):
    if not os.path.exists(template_path):
        return None
    try:
        doc = Document(template_path)
        formatted_date = l_date.strftime("%d-%m-%Y")
        for p in doc.paragraphs:
            if "[LetterDate]" in p.text:
                p.text = p.text.replace("[LetterDate]", formatted_date)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "[LetterDate]" in cell.text:
                        cell.text = cell.text.replace("[LetterDate]", formatted_date)
        create_noc_table(doc, emp_data_list)
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
    except Exception as e:
        st.error(f"Error: {e}")
        return None

# =================================================================
# --- 2. MAIN UI ---
# =================================================================
st.set_page_config(layout="wide", page_title="Railway NOC System")

if 'auth' not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    st.title("🔒 Login")
    with st.form("login"):
        u, p = st.text_input("User"), st.text_input("Pass", type="password")
        if st.form_submit_button("Login") and u == "admin" and p == "Sgam@4321":
            st.session_state.auth = True
            st.rerun()
    st.stop()

tab1, tab2 = st.tabs(["📝 Generate NOC", "📊 NOC Records Report"])

with tab1:
    st.header("Exam NOC Generator")
    df_emp = get_employees()
    if not df_emp.empty:
        df_emp['Display'] = df_emp['Employee Name'] + " (" + df_emp['PF Number'].astype(str) + ")"
        selected_display = st.multiselect("Karmchari Chunein", df_emp['Display'].tolist())
        final_list = []
        if selected_display:
            for display_val in selected_display:
                pf = display_val.split('(')[-1].replace(')', '').strip()
                row = df_emp[df_emp['PF Number'].astype(str) == pf].iloc[0]
                col1, col2 = st.columns([2, 1])
                with col1:
                    exam = st.text_input(f"Exam for {row['Employee Name']}", key=f"ex_{pf}")
                with col2:
                    count = get_noc_count_this_year(pf)
                    term = ["First", "Second", "Third", "Fourth"][count] if count < 4 else None
                    if term:
                        st.info(f"Term: {term}")
                        final_list.append({
                            "PFNumber": pf, "Name": row.get('Employee Name in Hindi', row['Employee Name']),
                            "Desig": row.get('Designation in Hindi', row['Designation']),
                            "ExamName": exam, "Term": term, "Year": datetime.now().year, "Timestamp": datetime.now()
                        })
                    else: st.error("Limit Exceeded")

            if st.button("Generate & Save") and final_list:
                out = generate_multi_noc(TEMPLATE_PATH, datetime.now(), final_list)
                if out:
                    for item in final_list: db.collection(NOC_HISTORY_COLLECTION).add(item)
                    st.success("Saved!")
                    st.download_button("📥 Download", out, "Joint_Exam_NOC_datetime.now().strftime('%d%m%Y').docx")

with tab2:
    st.header("📊 History Report")
    df_h = get_noc_history()
    if not df_h.empty:
        # Columns list jo hum dikhana chahte hain
        cols_to_show = ['Timestamp', 'PFNumber', 'Name', 'Designation', 'ExamName', 'Term', 'Year']
        # Sirf wahi columns select karein jo dataframe mein exist karte hain (Fix for KeyError)
        available_cols = [c for c in cols_to_show if c in df_h.columns]
        st.dataframe(df_h[available_cols], use_container_width=True)
    else:
        st.info("No records found.")
