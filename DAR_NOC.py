import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import io
import os
import firebase_admin
from firebase_admin import credentials, firestore

# --- 0. Configuration & Login ---
ADMIN_USER = "admin"
ADMIN_PASS = "Sgam@4321"

def check_login():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if not st.session_state.logged_in:
        st.title("🔐 Railway DAR System Login")
        with st.form("login_form"):
            user = st.text_input("Username")
            pas = st.text_input("Password", type="password")
            if st.form_submit_button("Login"):
                if user == ADMIN_USER and pas == ADMIN_PASS:
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("❌ Galat Username ya Password")
        return False
    return True

# --- 1. Firebase Initialization ---
@st.cache_resource
def init_db():
    if not firebase_admin._apps:
        try:
            if "firebase_config" in st.secrets:
                cred_dict = dict(st.secrets["firebase_config"])
                if isinstance(cred_dict.get('private_key'), str):
                    cred_dict['private_key'] = cred_dict['private_key'].replace('\\n', '\n')
                cred = credentials.Certificate(cred_dict)
            else:
                json_path = 'sgamoffice-firebase-adminsdk-fbsvc-253915b05b.json'
                cred = credentials.Certificate(json_path)
            firebase_admin.initialize_app(cred)
        except Exception as e:
            st.error(f"Firebase Error: {e}"); st.stop()
    return firestore.client()

# --- 2. Advanced Document Filling Logic ---
def fill_bulk_template(doc, selected_data, report_date):
    """Document ke har kone mein [Date] badalna aur aakhri table mein data bharna"""
    
    # A. Paragraphs mein date replace karein
    for p in doc.paragraphs:
        if "[Date]" in p.text:
            p.text = p.text.replace("[Date]", str(report_date))

    # B. Sabhi Tables mein date check karein (Header tables ke liye)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[Date]" in p.text:
                        p.text = p.text.replace("[Date]", str(report_date))

    # C. Headers/Footers mein date check karein
    for section in doc.sections:
        for header_p in section.header.paragraphs:
            if "[Date]" in header_p.text:
                header_p.text = header_p.text.replace("[Date]", str(report_date))

    # D. Data Table Filling (Sabse aakhri table)
    if not doc.tables:
        st.error("Template mein koi table nahi mili!")
        return doc
    
    data_table = doc.tables[-1] 
    num_cols = len(data_table.columns)

    for i, emp in enumerate(selected_data):
        # Index 1 pehli placeholder row hai
        if i == 0 and len(data_table.rows) >= 2:
            row_cells = data_table.rows[1].cells
        else:
            row_cells = data_table.add_row().cells
            
        # str() wrap kiya gaya hai taaki NoneType error na aaye
        if num_cols > 0: row_cells[0].text = str(i + 1)
        if num_cols > 1: row_cells[1].text = str(emp.get('name', ''))
        if num_cols > 2: row_cells[2].text = str(emp.get('desig', ''))
        if num_cols > 3: row_cells[3].text = str(emp.get('pf', ''))
        if num_cols > 4: row_cells[4].text = "कर्मचारी के विरूद्ध डी.ए.आर. एवं विजिलेंस केश लम्बित नहीं है"
            
    return doc

# --- 3. Main App Execution ---
def main():
    st.set_page_config(layout="wide", page_title="Railway DAR NOC")
    if not check_login(): return

    if st.sidebar.button("🚪 Logout"):
        st.session_state.logged_in = False
        st.rerun()

    db = init_db()
    
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    TEMPLATE_PATH = os.path.join(BASE_DIR, "assets", "DAR NOC temp.docx")

    # Employee list fetch
    docs = db.collection("employees").stream()
    df_emp = pd.DataFrame([{**d.to_dict(), 'id': d.id} for d in docs])

    tab1, tab2 = st.tabs(["📝 Generate Joint NOC", "📊 View Records"])

    with tab1:
        st.header("Bulk DAR NOC Generator")
        if not df_emp.empty:
            emp_options = df_emp.apply(lambda r: f"{r.get('Employee Name')} ({r.get('HRMS ID')})", axis=1).tolist()
            selected_names = st.multiselect("Staff Select Karein", emp_options)
            rep_date = st.date_input("Letter Date", value=datetime.now())

            if st.button("Generate Joint NOC Document"):
                if not selected_names:
                    st.warning("Kripya kam se kam ek staff select karein.")
                elif not os.path.exists(TEMPLATE_PATH):
                    st.error(f"Template File Missing: {TEMPLATE_PATH}")
                else:
                    # Selected Staff Details taiyar karein
                    selected_data = []
                    for name_str in selected_names:
                        h_id = name_str.split('(')[-1].strip(')')
                        row = df_emp[df_emp['HRMS ID'] == h_id].iloc[0]
                        selected_data.append({
                            "name": str(row.get('Employee Name', '')),
                            "desig": str(row.get('Designation', '')),
                            "pf": str(row.get('PF Number') if row.get('PF Number') else h_id)
                        })
                    
                    # Document Process karein
                    doc = Document(TEMPLATE_PATH)
                    formatted_date = rep_date.strftime("%d/%m/%Y")
                    filled_doc = fill_bulk_template(doc, selected_data, formatted_date)
                    
                    buf = io.BytesIO()
                    filled_doc.save(buf)
                    
                    # History Update
                    for emp in selected_data:
                        db.collection("dar_history").add({
                            "Employee Name": emp['name'],
                            "PF Number": emp['pf'],
                            "Generated_At": datetime.now(),
                            "Type": "Joint NOC"
                        })
                    
                    st.success(f"✅ Joint NOC Taiyar hai ({len(selected_data)} staff)!")
                    st.download_button(
                        label="📥 Download Joint NOC (.docx)",
                        data=buf.getvalue(),
                        file_name=f"DAR_NOC_{datetime.now().strftime('%d%m%Y')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning("Employees database mein nahi mile.")

    with tab2:
        st.header("📋 Generation History")
        h_docs = db.collection("dar_history").order_by("Generated_At", direction=firestore.Query.DESCENDING).limit(50).stream()
        h_list = [{**d.to_dict()} for d in h_docs]
        if h_list:
            st.dataframe(pd.DataFrame(h_list)[['Generated_At', 'Employee Name', 'PF Number']], use_container_width=True)

if __name__ == "__main__":
    main()
