import streamlit as st
import pandas as pd
import firebase_admin
from firebase_admin import credentials, firestore
import os
import io
from docx import Document
from datetime import date, datetime, timedelta

# --- 1. FIREBASE INITIALIZATION ---
if not firebase_admin._apps:
    try:
        cred_dict = dict(st.secrets["firebase_config"])
        if isinstance(cred_dict.get('private_key'), str):
            cred_dict['private_key'] = cred_dict['private_key'].replace('\\n', '\n')
        cred = credentials.Certificate(cred_dict)
        firebase_admin.initialize_app(cred)
    except Exception as e:
        st.error(f"Firebase Error: {e}")

db = firestore.client()

# --- 2. DATA CACHING ---
@st.cache_data(ttl=600)
def get_employee_data():
    try:
        emp_stream = db.collection('employees').stream()
        data = [d.to_dict() for d in emp_stream]
        return pd.DataFrame(data) if data else pd.DataFrame()
    except:
        return pd.DataFrame()

# --- 3. DOCUMENT ENGINE ---
def safe_replace(paragraph, context):
    inline = paragraph.runs
    if not inline: return
    full_text = "".join([run.text for run in inline])
    original_text = full_text
    for key, val in context.items():
        placeholder = f"[{key}]"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(val) if val is not None else "")
    if full_text != original_text:
        for i in range(len(inline)): inline[i].text = ""
        inline[0].text = full_text

def generate_doc(template_name, context):
    path = f"assets/{template_name}.docx"
    if not os.path.exists(path):
        st.error(f"Template not found: {path}")
        return None
    try:
        doc = Document(path)
        for p in doc.paragraphs: safe_replace(p, context)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: safe_replace(p, context)
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Doc Error: {e}")
        return None

# --- SESSION STATE INITIALIZATION (Buttons persistence) ---
if 'absent_docs' not in st.session_state:
    st.session_state.absent_docs = {"ready": False, "d_doc": None, "s_doc": None, "name": ""}
if 'sf11_doc' not in st.session_state:
    st.session_state.sf11_doc = {"ready": False, "doc": None, "name": ""}

# --- 4. MAIN INTERFACE ---
st.set_page_config(page_title="Railway Admin Pro", layout="wide")

pwd = st.sidebar.text_input("Password", type="password")
if pwd == st.secrets.get("PASSWORD", "sgam@4321"):
    
    with st.spinner("Loading Database..."):
        emp_df = get_employee_data()
        if not emp_df.empty:
            emp_df['Full_Disp'] = emp_df['PF No.'].astype(str) + " - " + emp_df['Employee Name in Hindi'].astype(str)

    tab = st.sidebar.radio("Navigation", ["Absent Case (Duty+SF11)", "Other SF-11/Order", "Appeal Process", "SF-11 Register & Import"])
    
    if st.sidebar.button("🔄 Refresh Data"):
        st.cache_data.clear()
        st.session_state.absent_docs = {"ready": False, "d_doc": None, "s_doc": None, "name": ""}
        st.session_state.sf11_doc = {"ready": False, "doc": None, "name": ""}
        st.rerun()

    LEGAL_ENDING = " जो कि रेल सेवक होने के नाते आपकी रेल सेवा निष्ठा के प्रति घोर लापरवाही को प्रदर्शित करता है। अतः आप कामों व भूलो के फेहरिस्त धारा 1, 2 एवं 3 के उल्लंघन के दोषी पाए जाते है।"

    # --- TAB 1: ABSENT CASE ---
    if tab == "Absent Case (Duty+SF11)":
        st.subheader("📝 अनुपस्थिति प्रकरण (Absent Case)")
        if not emp_df.empty:
            sel_emp = st.selectbox("कर्मचारी चुनें", emp_df['Full_Disp'].unique())
            r = emp_df[emp_df['Full_Disp'] == sel_emp].iloc[0]
            
            c1, c2 = st.columns(2)
            f_dt = c1.date_input("अनुपस्थिति से")
            t_dt = c2.date_input("अनुपस्थिति तक")
            
            if st.button("Generate Documents"):
                unit_2digit = str(r.get('Unit', ''))[:2]
                short_name = str(r.get('SF-11 short name', '')).strip()
                new_letter_no = f"सं/No./स्‍टॉफ/मानक फॉर्म/{short_name}/{unit_2digit}"
                
                ctx = {
                    "EmployeeName": r['Employee Name in Hindi'],
                    "Designation": r['Designation in Hindi'],
                    "ShortName": short_name,
                    "Unit": unit_2digit,
                    "PFNumber": str(r['PF No.']).strip(),
                    "FromDate": f_dt.strftime('%d-%m-%Y'),
                    "ToDate": t_dt.strftime('%d-%m-%Y'),
                    "DutyDate": (t_dt + timedelta(days=1)).strftime('%d-%m-%Y'),
                    "LetterDate": date.today().strftime('%d-%m-%Y'),
                    "LetterNo": new_letter_no,
                    "Memo": f"आप दिनांक {f_dt.strftime('%d-%m-%Y')} से {t_dt.strftime('%d-%m-%Y')} तक बिना सूचना अनुपस्थित रहे," + LEGAL_ENDING
                }
                st.session_state.absent_docs["d_doc"] = generate_doc("Absent Duty letter temp", ctx)
                st.session_state.absent_docs["s_doc"] = generate_doc("SF-11 temp", ctx)
                st.session_state.absent_docs["name"] = r['Employee Name in Hindi']
                st.session_state.absent_docs["ready"] = True
                
                db.collection("sf11_register").add({**ctx, "status": "Issued", "timestamp": datetime.now()})
                st.success("डेटाबेस में दर्ज और फाइलें तैयार!")

            if st.session_state.absent_docs["ready"]:
                d1, d2 = st.columns(2)
                if st.session_state.absent_docs["d_doc"]:
                    d1.download_button("⬇️ Duty Letter", st.session_state.absent_docs["d_doc"], f"Duty_{st.session_state.absent_docs['name']}.docx")
                if st.session_state.absent_docs["s_doc"]:
                    d2.download_button("⬇️ SF-11", st.session_state.absent_docs["s_doc"], f"SF11_{st.session_state.absent_docs['name']}.docx")

    # --- TAB 2: OTHER SF-11 & ORDER ---
    elif tab == "Other SF-11/Order":
        mode = st.radio("प्रकार", ["नया SF-11 जारी करें", "दण्‍डादेश (Punishment Order)"])
        
        if mode == "नया SF-11 जारी करें":
            st.subheader("📝 नया SF-11 जारी करें")
            if not emp_df.empty:
                sel_emp_sf = st.selectbox("कर्मचारी चुनें", emp_df['Full_Disp'].unique())
                r_sf = emp_df[emp_df['Full_Disp'] == sel_emp_sf].iloc[0]
                user_memo = st.text_area("आरोप का विवरण")
                
                if st.button("Generate SF-11"):
                    unit_2digit = str(r_sf.get('Unit', ''))[:2]
                    short_name = str(r_sf.get('SF-11 short name', '')).strip()
                    new_letter_no = f"सं/No./स्‍टॉफ/मानक फॉर्म/{short_name}/{unit_2digit}"
                    
                    ctx = {
                        "EmployeeName": r_sf['Employee Name in Hindi'],
                        "Designation": r_sf['Designation in Hindi'],
                        "ShortName": short_name,
                        "Unit": unit_2digit,
                        "PFNumber": str(r_sf['PF No.']).strip(),
                        "LetterDate": date.today().strftime('%d-%m-%Y'),
                        "Memo": user_memo + LEGAL_ENDING,
                        "LetterNo": new_letter_no
                    }
                    st.session_state.sf11_doc["doc"] = generate_doc("SF-11 temp", ctx)
                    st.session_state.sf11_doc["name"] = r_sf['Employee Name in Hindi']
                    st.session_state.sf11_doc["ready"] = True
                    db.collection("sf11_register").add({**ctx, "status": "Issued", "timestamp": datetime.now()})
                    st.success("SF-11 तैयार!")

                if st.session_state.sf11_doc["ready"]:
                    st.download_button("⬇️ Download SF-11", st.session_state.sf11_doc["doc"], f"SF11_{st.session_state.sf11_doc['name']}.docx")

        elif mode == "दण्‍डादेश (Punishment Order)":
            st.subheader("🔨 दण्‍डादेश (NIP) जनरेट")
            docs = db.collection("sf11_register").where("status", "==", "Issued").stream()
            reg_list = [d.to_dict() | {"doc_id": d.id} for d in docs]
            
            if reg_list:
                reg_df = pd.DataFrame(reg_list)
                reg_df['Select_Disp'] = reg_df['PFNumber'].astype(str) + " - " + reg_df['EmployeeName'] + " - Dt: " + reg_df['LetterDate']
                case = reg_df[reg_df['Select_Disp'] == st.selectbox("केस चुनें", reg_df['Select_Disp'].unique())].iloc[0]
                
                unit_extracted = ""
                if not emp_df.empty:
                    match = emp_df[emp_df['PF No.'].astype(str).str.strip() == str(case['PFNumber']).strip()]
                    if not match.empty:
                        unit_extracted = str(match.iloc[0].get('Unit', ''))[:2]
                
                c1, c2 = st.columns(2)
                dandadesh_no = c1.text_input("दण्‍डादेश क्रमांक", value=f"SGAM/NIP/{case['LetterNo']}")
                order_date = c2.date_input("दण्‍डादेश दिनांक", value=date.today())
                punishment_text = st.selectbox("दण्ड चुनें", ["आगामी देय एक वर्ष की वेतन वृद्धि असंचयी प्रभाव से अवरोधित किए जाने की शास्ति दी जाती है।", "आगामी देय एक वर्ष की वेतन वृद्धि संचयी प्रभाव से अवरोधित किए जाने की शास्ति दी जाती है।", "आगामी देय एक सेट सुविधा पास अवरोधित किए जाने की शास्ति दी जाती है।" , "आगामी देय एक सेट पीटीओ अवरोधित किए जाने की शास्ति दी जाती है।" ])

                if st.button("Generate & Update"):
                    ctx = {
                        "EmployeeName": case.get('EmployeeName', ''),
                        "Designation": case.get('Designation', ''),
                        "Unit": unit_extracted,
                        "Dandadesh": dandadesh_no,
                        "LetterNo.": case.get('LetterNo', ''),
                        "SF-11Date": case.get('LetterDate', ''),
                        "LetterDate": order_date.strftime('%d-%m-%Y'),
                        "OrderDate": order_date.strftime('%d-%m-%Y'),
                        "Memo": punishment_text
                    }
                    doc_bio = generate_doc("SF-11 Punishment order temp", ctx)
                    if doc_bio:
                        db.collection("sf11_register").document(case['doc_id']).update({
                            "OrderNo": dandadesh_no, "PunishmentDetails": punishment_text, "OrderDate": order_date.strftime('%d-%m-%Y'),
                            "status": "Closed"
                        })
                        st.download_button("⬇️ Download NIP", doc_bio, f"NIP_{case['EmployeeName']}.docx")
            else: st.warning("कोई पेंडिंग केस नहीं मिला।")

    # --- TAB 3: APPEAL PROCESS ---
    elif tab == "Appeal Process":
        st.subheader("⚖️ अपील प्रबंधन")
        appeal_mode = st.radio("चुनें", ["1. अपील दर्ज करें", "2. अपील निर्णय"])

        if appeal_mode == "1. अपील दर्ज करें":
            docs = db.collection("sf11_register").where("status", "==", "Closed").stream()
            closed_list = [d.to_dict() | {"doc_id": d.id} for d in docs]
            if closed_list:
                df_c = pd.DataFrame(closed_list)
                sel = st.selectbox("अपील हेतु केस चुनें", df_c['EmployeeName'].unique())
                case = df_c[df_c['EmployeeName'] == sel].iloc[0]
                app_date = st.date_input("अपील प्राप्ति दिनांक")
                
                if st.button("Generate Appeal Letter"):
                    ctx = {**case, "AppealDate": app_date.strftime('%d-%m-%Y')}
                    doc = generate_doc("apeal_letter_temp", ctx)
                    if doc:
                        db.collection("sf11_register").document(case['doc_id']).update({"status": "Appeal-Process", "AppealDate": ctx["AppealDate"]})
                        st.download_button("⬇️ Download Appeal Letter", doc, f"Appeal_{case['EmployeeName']}.docx")
            else: st.info("कोई क्लोज्ड केस नहीं मिला।")

        elif appeal_mode == "2. अपील निर्णय":
            docs = db.collection("sf11_register").where("status", "==", "Appeal-Process").stream()
            proc_list = [d.to_dict() | {"doc_id": d.id} for d in docs]
            if proc_list:
                df_p = pd.DataFrame(proc_list)
                sel = st.selectbox("निर्णय हेतु केस चुनें", df_p['EmployeeName'].unique())
                case = df_p[df_p['EmployeeName'] == sel].iloc[0]
                decision = st.selectbox("निर्णय", ["दण्ड यथावत", "दण्ड कम", "दण्ड रद्द"])
                if st.button("Finalize"):
                    db.collection("sf11_register").document(case['doc_id']).update({"status": "Appeal-Closed", "AppealDecision": decision})
                    st.success("अपील क्लोज कर दी गई।")

    # --- TAB 4: REGISTER ---
    elif tab == "SF-11 Register & Import":
        st.subheader("📊 मास्टर रजिस्टर")
        if st.button("Load All Records"):
            all_docs = db.collection("sf11_register").get()
            all_reg = [d.to_dict() for d in all_docs]
            if all_reg:
                df_final = pd.DataFrame(all_reg)
                st.dataframe(df_final, use_container_width=True)
                csv = df_final.to_csv(index=False).encode('utf-8-sig')
                st.download_button("💾 Download CSV", csv, "Register.csv", "text/csv")
            else: st.warning("डेटा उपलब्ध नहीं है।")

else:
    st.info("Side menu में पासवर्ड डालें।")
